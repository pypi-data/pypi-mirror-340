from typing import Dict, Any, Optional, List
import re
from pathlib import Path
import logging
from PIL import Image
import pytesseract
import fitz  # PyMuPDF
import docx
import os

logger = logging.getLogger(__name__)

class DocumentPreprocessor:
    """Preprocesses documents to ensure quality before analysis."""
    
    MIN_TEXT_LENGTH = 50  # Minimum characters for meaningful content
    MIN_IMAGE_SIZE = 300  # Minimum dimension in pixels for reliable OCR
    
    def __init__(self):
        self.min_text_length = self.MIN_TEXT_LENGTH
        self.max_text_length = 100000  # Maximum characters for Gemini
    
    def preprocess(self, file_path: Path, mime_type: str) -> Dict[str, Any]:
        """Preprocess the document and return analysis results."""
        try:
            # Extract text based on file type
            if mime_type == 'application/pdf':
                result = self._extract_pdf_text(file_path)
            elif mime_type == 'application/vnd.openxmlforms-officedocument.wordprocessingml.document':
                result = self._extract_docx_text(file_path)
            elif mime_type == 'text/plain':
                result = self._extract_txt_text(file_path)
            elif mime_type in ['image/jpeg', 'image/png', 'image/webp']:
                # Images are handled directly by Gemini Vision
                return {
                    "status": "success",
                    "text": "",
                    "is_ready_for_gemini": True,
                    "analysis": {"type": "image"},
                    "recommendations": ["Use Gemini Vision for analysis"]
                }
            else:
                raise ValueError(f"Unsupported MIME type: {mime_type}")
            
            # Analyze the extracted text
            analysis = self._analyze_text(result["text"])
            
            # Check if document is ready for Gemini
            is_ready = self._check_ready_for_gemini(result, analysis)
            
            return {
                "status": "success",
                "text": result["text"],
                "is_ready_for_gemini": is_ready,
                "analysis": analysis,
                "recommendations": self._generate_recommendations(result, analysis, is_ready)
            }
            
        except Exception as e:
            logger.error(f"Error preprocessing document: {str(e)}")
            return {
                "status": "error",
                "error": str(e),
                "text": "",
                "is_ready_for_gemini": False,
                "analysis": None,
                "recommendations": ["Fix the error before proceeding with analysis"]
            }
    
    def _extract_pdf_text(self, file_path: Path) -> Dict[str, Any]:
        """Extract text from PDF file."""
        logger.info(f"Starting PDF text extraction for file: {file_path}")
        doc = None
        try:
            # Open the PDF document
            doc = fitz.open(str(file_path))
            logger.info(f"PDF opened successfully. Page count: {len(doc)}")
            
            if len(doc) == 0:
                logger.error("PDF is empty")
                raise ValueError("PDF is empty")
            
            text = ""
            needs_ocr = []
            total_pages = len(doc)
            
            for page_num in range(total_pages):
                try:
                    page = doc[page_num]
                    logger.debug(f"Processing page {page_num + 1}/{total_pages}")
                    
                    # Try different text extraction methods
                    page_text = None
                    
                    # Method 1: Try standard text extraction
                    page_text = page.get_text("text")
                    if not page_text.strip():
                        logger.warning(f"No text found on page {page_num + 1} using standard extraction")
                        
                        # Method 2: Try with layout preservation
                        page_text = page.get_text("blocks")
                        if not page_text.strip():
                            logger.warning(f"No text found on page {page_num + 1} using block extraction")
                            
                            # Method 3: Try with HTML output
                            page_text = page.get_text("html")
                            if not page_text.strip():
                                logger.warning(f"No text found on page {page_num + 1} using HTML extraction")
                    
                    # If still no text, mark for OCR
                    if not page_text.strip():
                        logger.warning(f"No text could be extracted from page {page_num + 1}, marking for OCR")
                        needs_ocr.append(page_num + 1)
                        page_text = f"\n[OCR needed for page {page_num + 1}]\n"
                    else:
                        logger.info(f"Successfully extracted text from page {page_num + 1}")
                    
                    text += page_text + "\n\n"  # Add spacing between pages
                    
                except Exception as page_error:
                    logger.error(f"Error processing page {page_num + 1}: {str(page_error)}")
                    needs_ocr.append(page_num + 1)
                    text += f"\n[Error processing page {page_num + 1}]\n"
            
            # Validate final text
            if not text.strip() and not needs_ocr:
                logger.error("No text could be extracted from the PDF")
                raise ValueError("No text could be extracted from the PDF")
            
            logger.info(f"PDF processing completed. Text length: {len(text)} characters")
            if needs_ocr:
                logger.warning(f"Pages requiring OCR: {needs_ocr}")
            
            return {
                "text": text,
                "needs_ocr": needs_ocr,
                "page_count": total_pages,
                "successful_pages": total_pages - len(needs_ocr)
            }
            
        except Exception as e:
            logger.error(f"Error in PDF text extraction: {str(e)}")
            raise
            
        finally:
            if doc:
                doc.close()
                logger.debug("PDF document closed")
    
    def _extract_docx_text(self, file_path: Path) -> Dict[str, Any]:
        """Extract text from DOCX file."""
        doc = docx.Document(file_path)
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        
        return {
            "text": text,
            "paragraph_count": len(doc.paragraphs)
        }
    
    def _extract_txt_text(self, file_path: Path) -> Dict[str, Any]:
        """Extract text from TXT file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
        except UnicodeDecodeError:
            # If UTF-8 fails, try reading as binary
            with open(file_path, 'rb') as f:
                text = f.read().decode('latin-1')  # Fallback to latin-1
        
        return {
            "text": text,
            "encoding": "utf-8" if 'encoding' not in locals() else "latin-1"
        }
    
    def _extract_image_text(self, file_path: Path) -> Dict[str, Any]:
        """Extract text from image using OCR."""
        try:
            # Open and check image
            image = Image.open(file_path)
            width, height = image.size
            
            if width < self.MIN_IMAGE_SIZE or height < self.MIN_IMAGE_SIZE:
                raise ValueError(f"Image too small for reliable OCR. Minimum size: {self.MIN_IMAGE_SIZE}x{self.MIN_IMAGE_SIZE}")
            
            # Convert to grayscale for better OCR
            if image.mode != 'L':
                image = image.convert('L')
            
            # Perform OCR
            text = pytesseract.image_to_string(image)
            
            return {
                "text": text,
                "dimensions": {"width": width, "height": height},
                "ocr_quality": "good" if text.strip() else "poor"
            }
            
        except Exception as e:
            logger.error(f"Error performing OCR: {str(e)}")
            raise
    
    def _analyze_text(self, text: str) -> Dict[str, Any]:
        """Analyze the extracted text for content quality and characteristics."""
        # Basic content analysis
        char_count = len(text)
        word_count = len(text.split())
        line_count = len(text.splitlines())
        
        # Content type detection
        has_numbers = bool(re.search(r'\d', text))
        has_dates = bool(re.search(r'\d{1,2}[/-]\d{1,2}[/-]\d{2,4}|\d{4}[/-]\d{1,2}[/-]\d{1,2}', text))
        has_emails = bool(re.search(r'[\w\.-]+@[\w\.-]+\.\w+', text))
        has_urls = bool(re.search(r'https?://\S+', text))
        
        # Content structure analysis
        paragraphs = text.split('\n\n')
        sentences = re.split(r'[.!?]+', text)
        
        # Check for potential data tables
        has_tables = bool(re.search(r'\|\s*[\w\s]+\s*\|', text)) or bool(re.search(r'\+[-]+\+', text))
        
        # Check for potential lists
        has_lists = bool(re.search(r'^\s*[\d\w]\.\s|\s*[-*]\s', text, re.MULTILINE))
        
        # Check for potential headers/sections
        has_headers = bool(re.search(r'^\s*[A-Z][A-Z\s]+\s*$', text, re.MULTILINE))
        
        # Check for potential code blocks
        has_code = bool(re.search(r'```|`[^`]+`', text))
        
        return {
            "content_metrics": {
                "char_count": char_count,
                "word_count": word_count,
                "line_count": line_count,
                "paragraph_count": len(paragraphs),
                "sentence_count": len(sentences),
                "avg_words_per_sentence": word_count / max(1, len(sentences)),
                "avg_words_per_paragraph": word_count / max(1, len(paragraphs))
            },
            "content_features": {
                "has_numbers": has_numbers,
                "has_dates": has_dates,
                "has_emails": has_emails,
                "has_urls": has_urls,
                "has_tables": has_tables,
                "has_lists": has_lists,
                "has_headers": has_headers,
                "has_code": has_code
            },
            "content_quality": {
                "is_too_short": char_count < self.min_text_length,
                "has_meaningful_content": word_count > 10 and len(paragraphs) > 1,
                "has_structure": has_headers or has_lists or has_tables
            }
        }
    
    def _check_ready_for_gemini(self, result: Dict[str, Any], analysis: Dict[str, Any]) -> bool:
        """Check if the document is ready for Gemini analysis."""
        # Check if text is too short
        if analysis["content_quality"]["is_too_short"]:
            return False
        
        # Check if content is meaningful
        if not analysis["content_quality"]["has_meaningful_content"]:
            return False
        
        # Check if PDF needs OCR
        if "needs_ocr" in result and result["needs_ocr"]:
            return False
        
        # Check if image OCR was successful
        if "ocr_quality" in result and result["ocr_quality"] == "poor":
            return False
        
        return True
    
    def _generate_recommendations(self, result: Dict[str, Any], analysis: Dict[str, Any], is_ready: bool) -> List[str]:
        """Generate recommendations based on content analysis."""
        recommendations = []
        
        if not is_ready:
            if analysis["content_quality"]["is_too_short"]:
                recommendations.append("Document content is too short for meaningful analysis")
            if not analysis["content_quality"]["has_meaningful_content"]:
                recommendations.append("Document appears to lack meaningful content")
            if "needs_ocr" in result and result["needs_ocr"]:
                recommendations.append(f"OCR needed for pages: {', '.join(map(str, result['needs_ocr']))}")
            if "ocr_quality" in result and result["ocr_quality"] == "poor":
                recommendations.append("Image quality may be too poor for reliable OCR")
        
        # Content-specific recommendations
        if analysis["content_features"]["has_tables"]:
            recommendations.append("Document contains tables - ensure proper formatting for analysis")
        if analysis["content_features"]["has_code"]:
            recommendations.append("Document contains code blocks - ensure proper syntax highlighting")
        if analysis["content_features"]["has_urls"]:
            recommendations.append("Document contains URLs - consider checking link validity")
        
        return recommendations
    
    def _analyze_structure(self, text: str) -> float:
        """Analyze document structure quality."""
        score = 0.0
        
        # Check for paragraphs
        paragraphs = text.split('\n\n')
        if len(paragraphs) > 1:
            score += 0.3
        
        # Check for sentences
        sentences = re.split(r'[.!?]+', text)
        if len(sentences) > 1:
            score += 0.3
        
        # Check for common formatting
        if any(char in text for char in ['\n', '\t', '  ']):
            score += 0.2
        
        # Check for common document elements
        if any(marker in text for marker in ['Chapter', 'Section', 'Introduction', 'Conclusion']):
            score += 0.2
        
        return min(score, 1.0) 