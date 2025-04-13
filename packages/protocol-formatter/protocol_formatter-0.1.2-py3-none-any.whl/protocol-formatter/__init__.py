import os
import re
import subprocess
import sys
from datetime import datetime

os.environ["DYLD_LIBRARY_PATH"] = "/opt/homebrew/lib"
os.environ["PKG_CONFIG_PATH"] = "/opt/homebrew/lib/pkgconfig"

import svgwrite
import cairosvg
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls


# Get the current date and time
def timestamp():
    """
    Returns the current date and time formatted as a string.
    """
    return datetime.now().strftime("%m-%d %H:%M")

def substitute_text(text):
    # Substitute temperature in Celsius
    text = re.sub(r'(\d+)C', r'\1°C', text)

    # Substitute Greek letters
    greek_letters = {
        'alpha': 'α', 'beta': 'β', 'gamma': 'γ', 'delta': 'δ', 'epsilon': 'ε',
        'zeta': 'ζ', 'eta': 'η', 'theta': 'θ', 'iota': 'ι', 'kappa': 'κ',
        'lambda': 'λ', 'mu': 'μ', 'nu': 'ν', 'xi': 'ξ', 'omicron': 'ο',
        'pi': 'π', 'rho': 'ρ', 'sigma': 'σ', 'tau': 'τ', 'upsilon': 'υ',
        'phi': 'φ', 'chi': 'χ', 'psi': 'ψ', 'omega': 'ω',
        'upperalpha': 'Α', 'upperbeta': 'Β', 'uppergamma': 'Γ', 'upperdelta': 'Δ',
        'upperepsilon': 'Ε', 'upperzeta': 'Ζ', 'uppereta': 'Η', 'uppertheta': 'Θ',
        'upperiota': 'Ι', 'upperkappa': 'Κ', 'upperlambda': 'Λ', 'uppermu': 'Μ',
        'uppernu': 'Ν', 'upperxi': 'Ξ', 'upperomicron': 'Ο', 'upperpi': 'Π',
        'upperrho': 'Ρ', 'uppersigma': 'Σ', 'uppertau': 'Τ', 'upperupsilon': 'Υ',
        'upperphi': 'Φ', 'upperchi': 'Χ', 'upperpsi': 'Ψ', 'upperomega': 'Ω'
    }
    for word, symbol in greek_letters.items():
        text = re.sub(r'\b' + word + r'\b', symbol, text)

    # Substitute micro- units
    micro_units = {
        'microLitre': 'μL', 'microgram': 'μg', 'micromol': 'μM'
    }
    for word, symbol in micro_units.items():
        text = re.sub(r'\b' + word + r'\b', symbol, text)

    return text

def create_arrow_svg(output_dir):
    svg_path = os.path.join(output_dir, 'arrow.svg')
    png_path = os.path.join(output_dir, 'arrow.png')

    dwg = svgwrite.Drawing(filename=svg_path, size=(20, 20))
    dwg.add(dwg.polygon(points=[(10, 20), (20, 10), (15, 10), (15, 0), (5, 0), (5, 10), (0, 10)], fill='black'))
    dwg.save()

    # Convert SVG to PNG using a file object
    with open(svg_path, 'rb') as svg_file, open(png_path, 'wb') as png_file:
        cairosvg.svg2png(file_obj=svg_file, write_to=png_file)

    return svg_path, png_path

def add_custom_styles(doc):
    styles = doc.styles

    # Create a custom style for first-level bullet points
    if 'CustomBullet' not in styles:
        style = styles.add_style('CustomBullet', WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = 'Calibri'
        style.font.size = Pt(11)
        paragraph_format = style.paragraph_format
        paragraph_format.left_indent = Inches(0.25)
        paragraph_format.first_line_indent = Inches(-0.25)
        paragraph_format.space_before = Pt(0)
        paragraph_format.space_after = Pt(0)

    # Create a custom style for second-level bullet points
    if 'CustomBullet2' not in styles:
        style = styles.add_style('CustomBullet2', WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = 'Calibri'
        style.font.size = Pt(11)
        paragraph_format = style.paragraph_format
        paragraph_format.left_indent = Inches(0.5)
        paragraph_format.first_line_indent = Inches(-0.25)
        paragraph_format.space_before = Pt(0)
        paragraph_format.space_after = Pt(0)

def add_numbering_part(doc):
    numbering = doc.part.numbering_part.numbering_definitions._numbering
    abstract_num_id = 0  # Use 0 if no other numbering is defined

    abstractNum = OxmlElement('w:abstractNum')
    abstractNum.set(qn('w:abstractNumId'), str(abstract_num_id))

    # First-level bullet
    lvl0 = OxmlElement('w:lvl')
    lvl0.set(qn('w:ilvl'), '0')

    start = OxmlElement('w:start')
    start.set(qn('w:val'), '1')
    lvl0.append(start)

    numFmt = OxmlElement('w:numFmt')
    numFmt.set(qn('w:val'), 'bullet')
    lvl0.append(numFmt)

    lvlText = OxmlElement('w:lvlText')
    lvlText.set(qn('w:val'), '•')  # Filled circle
    lvl0.append(lvlText)

    lvlJc = OxmlElement('w:lvlJc')
    lvlJc.set(qn('w:val'), 'left')
    lvl0.append(lvlJc)

    pPr = OxmlElement('w:pPr')
    ind = OxmlElement('w:ind')
    ind.set(qn('w:left'), '720')  # 0.5 inches
    ind.set(qn('w:hanging'), '360')  # 0.25 inches
    pPr.append(ind)
    lvl0.append(pPr)

    rPr = OxmlElement('w:rPr')
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), 'Symbol')  # Font for '•'
    rFonts.set(qn('w:hAnsi'), 'Symbol')
    rPr.append(rFonts)
    lvl0.append(rPr)

    abstractNum.append(lvl0)

    # Second-level bullet
    lvl1 = OxmlElement('w:lvl')
    lvl1.set(qn('w:ilvl'), '1')

    start = OxmlElement('w:start')
    start.set(qn('w:val'), '1')
    lvl1.append(start)

    numFmt = OxmlElement('w:numFmt')
    numFmt.set(qn('w:val'), 'bullet')  # Set to 'bullet'
    lvl1.append(numFmt)

    lvlText = OxmlElement('w:lvlText')
    lvlText.set(qn('w:val'), '○')  # Unfilled circle
    lvl1.append(lvlText)

    lvlJc = OxmlElement('w:lvlJc')
    lvlJc.set(qn('w:val'), 'left')
    lvl1.append(lvlJc)

    pPr = OxmlElement('w:pPr')
    ind = OxmlElement('w:ind')
    ind.set(qn('w:left'), '1080')  # 0.75 inches
    ind.set(qn('w:hanging'), '360')  # 0.25 inches
    pPr.append(ind)
    lvl1.append(pPr)

    rPr = OxmlElement('w:rPr')
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), 'Calibri')  # Font for '○'
    rFonts.set(qn('w:hAnsi'), 'Calibri')
    rPr.append(rFonts)
    lvl1.append(rPr)

    abstractNum.append(lvl1)

    numbering.append(abstractNum)

    num = OxmlElement('w:num')
    num.set(qn('w:numId'), '1')
    abstractNumId_elem = OxmlElement('w:abstractNumId')
    abstractNumId_elem.set(qn('w:val'), str(abstract_num_id))
    num.append(abstractNumId_elem)
    numbering.append(num)

def create_protocol_doc(source, title, steps, output_dir, output_filename, arrow_image_path):
    # Create a Word document
    doc = Document()
    add_custom_styles(doc)
    add_numbering_part(doc)

    # Add date and protocol source to header
    section = doc.sections[0]
    header = section.header
    header_table = header.add_table(rows=1, cols=2, width=Inches(6))
    header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    header_cells = header_table.rows[0].cells
    today_date = datetime.today().strftime('%Y%m%d')
    header_cells[0].text = f'Protocol prep date: {today_date}'
    header_cells[0].paragraphs[0].runs[0].font.size = Pt(9)
    header_cells[0].paragraphs[0].runs[0].italic = True
    header_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    header_cells[1].text = f'Protocol Source: {source}'
    header_cells[1].paragraphs[0].runs[0].font.size = Pt(9)
    header_cells[1].paragraphs[0].runs[0].italic = True
    header_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Add protocol title
    title_paragraph = doc.add_paragraph()
    title_run = title_paragraph.add_run(format(title))
    title_run.bold = True
    title_run.underline = True
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add protocol steps using tables to encapsulate multiple paragraphs
    for i, step in enumerate(steps):
        lines = step.split('\n')

        # Create a table with 1 row and 1 column for the step
        table = doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        table.columns[0].width = Inches(6)  # Adjust width as needed

        cell = table.cell(0, 0)

        # Set cell margins to zero
        cell.margin_top = Cm(0)
        cell.margin_bottom = Cm(0)
        cell.margin_left = Cm(0)
        cell.margin_right = Cm(0)

        # Set cell shading
        tcPr = cell._tc.get_or_add_tcPr()
        shading_xml = r'<w:shd {} w:val="clear" w:color="auto" w:fill="d4f4fa"/>'.format(nsdecls('w'))
        shading_element = parse_xml(shading_xml)
        tcPr.append(shading_element)

        # Set cell borders
        tcBorders = OxmlElement('w:tcBorders')
        for border_name in ('top', 'left', 'bottom', 'right'):
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '8')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), 'auto')
            tcBorders.append(border)
        tcPr.append(tcBorders)

        # Now, for each line, add a paragraph to the cell
        for idx, line in enumerate(lines):
            line = substitute_text(line)  # Apply text substitutions
            if idx == 0:
                # First line: step title
                p = cell.add_paragraph()
                if ': ' in line:
                    before_colon, after_colon = line.split(': ', 1)
                    run_bold = p.add_run(before_colon + ': ')
                    run_bold.bold = True
                    run_normal = p.add_run(after_colon)
                else:
                    run_bold = p.add_run(line)
                    run_bold.bold = True
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                # Set font properties for both runs
                for run in p.runs:
                    run.font.size = Pt(11)
                    run.font.name = 'Calibri'
                    run.font.color.rgb = RGBColor(0, 0, 0)
            else:
                if line.startswith('* '):
                    # First-level bullet point
                    text = line[2:].strip()
                    p = cell.add_paragraph(style='CustomBullet')
                    run = p.add_run(text)
                    # Assign numbering level and ID
                    numPr = OxmlElement('w:numPr')
                    ilvl = OxmlElement('w:ilvl')
                    ilvl.set(qn('w:val'), '0')
                    numId = OxmlElement('w:numId')
                    numId.set(qn('w:val'), '1')
                    numPr.append(ilvl)
                    numPr.append(numId)
                    p._p.get_or_add_pPr().append(numPr)
                elif line.startswith('** '):
                    # Second-level bullet point
                    text = line[3:].strip()
                    p = cell.add_paragraph(style='CustomBullet2')
                    run = p.add_run(text)
                    # Assign numbering level and ID
                    numPr = OxmlElement('w:numPr')
                    ilvl = OxmlElement('w:ilvl')
                    ilvl.set(qn('w:val'), '1')
                    numId = OxmlElement('w:numId')
                    numId.set(qn('w:val'), '1')
                    numPr.append(ilvl)
                    numPr.append(numId)
                    p._p.get_or_add_pPr().append(numPr)
                else:
                    # Normal text
                    p = cell.add_paragraph()
                    run = p.add_run(line)
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                # Set font properties
                run.font.size = Pt(11)
                run.font.name = 'Calibri'
                run.font.color.rgb = RGBColor(0, 0, 0)

        # Add arrow below the bordered box, centered
        if i < len(steps) - 1:
            arrow_paragraph = doc.add_paragraph()
            arrow_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            arrow_paragraph.paragraph_format.space_before = Pt(0)
            arrow_paragraph.paragraph_format.space_after = Pt(0)
            run = arrow_paragraph.add_run()
            run.add_picture(arrow_image_path, width=Inches(0.3), height=Inches(0.3))

    # Add page number to footer
    footer = section.footer
    footer.is_linked_to_previous = False  # Ensure footer is not linked to previous sections
    footer_paragraph = footer.paragraphs[0]
    footer_paragraph.text = "Page "
    footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer_paragraph.add_run()
    fldChar = OxmlElement('w:fldChar')  # creates a new element
    fldChar.set(qn('w:fldCharType'), 'begin')  # sets attribute on element
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')  # sets attribute on element
    instrText.text = "PAGE"
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')  # sets attribute on element
    run._r.append(fldChar)
    run._r.append(instrText)
    run._r.append(fldChar2)

    # Save the Word document
    docx_path = os.path.join(output_dir, f"{output_filename}.docx")
    doc.save(docx_path)
    print(f"[{timestamp()}] Word document saved at {docx_path}")

    return docx_path

def convert_docx_to_pdf(docx_path, output_dir, output_filename):
    pdf_path = os.path.join(output_dir, f"{output_filename}.pdf")

    # Convert the DOCX file to PDF using LibreOffice
    try:
        print()
        subprocess.run(
            ['/Applications/LibreOffice.app/Contents/MacOS/soffice', '--headless', '--convert-to', 'pdf', '--outdir',
             output_dir, docx_path], check=True)
        print()
        print(f"[{timestamp()}] PDF document saved at {pdf_path}")
    except subprocess.CalledProcessError as e:
        print(f"[{timestamp()}] Error converting DOCX to PDF: {e}")

def get_protocol_info():
    source = input(f"[{timestamp()}] Enter the protocol source: ")
    title = input(f"[{timestamp()}] Enter the protocol title: ")
    steps = []
    print(f"[{timestamp()}] Enter the protocol steps (type 'done' to finish):")
    while True:
        step = input()
        if step.lower() == 'done':
            break
        steps.append(step)
    return source, title, steps

def read_protocol_from_file(file_path):
    with open(file_path, 'r') as file:
        lines = file.readlines()

    source = lines[0].strip()[lines[0].find(":") + 1:].strip()
    title = lines[1].strip()[lines[1].find(":") + 1:].strip()

    steps = []
    current_step = ""
    in_step = False

    for line in lines[2:]:
        stripped_line = line.rstrip('\n')
        if stripped_line.startswith("-"):
            # If we are already in a step, add the current step to the steps list
            if in_step:
                steps.append(current_step.strip())
                current_step = ""
            in_step = True
            # Start the new step without the dash
            current_step += stripped_line[1:].strip() + "\n"
        elif in_step:
            # Handle bullet points and indented bullet points within a step
            if stripped_line.startswith("* "):
                current_step += "* " + stripped_line[2:].strip() + "\n"
            elif stripped_line.startswith("** "):
                current_step += "** " + stripped_line[3:].strip() + "\n"
            else:
                # Continue adding to the current step
                current_step += stripped_line + "\n"

    # Add the last step if it exists
    if current_step:
        steps.append(current_step.strip())

    return source, title, steps

def main():
    mode = input(f"[{timestamp()}] Do you want to enter protocol info manually or from a file? (manual/file): ").strip().lower()
    if mode == 'file':
        file_path = input(f"[{timestamp()}] Enter the path to the text file: ").strip()
        source, title, steps = read_protocol_from_file(file_path)
        output_filename = os.path.basename(file_path).split(".txt")[0]
        pre_output_dir = os.path.dirname(os.path.dirname(file_path))
    else:
        source, title, steps = get_protocol_info()
        pre_output_dir = os.getcwd()
        output_filename = 'protocol'

    output_dir = input(f"[{timestamp()}] Take output directory as '{pre_output_dir}'?\n"
                       f"              If you'd like to define output directory, please enter its path. Else simply press enter: ").strip()
    if output_dir == "" or not os.path.exists(output_dir):
        print(f"[{timestamp()}] No valid path input found taking '{pre_output_dir}' as output directory.")
        output_dir = pre_output_dir

    svg_path, png_path = create_arrow_svg(output_dir)  # Create the arrow SVG and convert to PNG
    docx_path = create_protocol_doc(source, title, steps, output_dir, output_filename, png_path)
    convert_docx_to_pdf(docx_path, output_dir, output_filename)

    # Delete the temporary arrow files
    try:
        os.remove(svg_path)
        os.remove(png_path)
        print(f"[{timestamp()}] Temporary arrow files deleted: {svg_path}, {png_path}")
    except OSError as e:
        print(f"[{timestamp()}] Error deleting temporary arrow files: {e}")

if __name__ == "__main__":
    main()
