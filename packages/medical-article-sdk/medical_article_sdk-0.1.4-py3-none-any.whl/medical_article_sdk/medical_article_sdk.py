#!/usr/bin/env python
# -*- coding: utf-8 -*-

"""
医学讲座转写稿批量处理工具 - 主程序

本程序用于批量处理医学讲座转写稿，将其转化为结构清晰的医学文章，
并自动生成三种风格的标题，最后保存为Word和Markdown格式。

作者: AI助手
日期: 2025-04-04
"""

import json
import os
import logging
import requests
import time
import argparse
import sys
import re
import traceback
from datetime import datetime
from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
import importlib.resources as resources

class MedicalArticleProcessor:
    """医学讲座转写稿处理器

    用于批量处理医学讲座转写稿，生成规范医学文章及标题。
    支持生成不同风格的标题，保存为多种格式，并生成处理报告。
    """

    def __init__(self, api_key,config_path='config.json', retry_only=False):
        """初始化处理器

        Args:
            config_path (str): 配置文件路径
            retry_only (bool): 是否只处理之前失败的文件
        """
        self.retry_only = retry_only
        self.api_key = api_key
        # 当前处理时间，用于报告
        self.current_time = datetime.now().strftime('%Y%m%d_%H%M%S')
        self.config = self.load_config(config_path)
        self.ensure_directories()
        self.setup_logging()
        self.medical_prompt = self.load_prompt(self.config['prompt_path'])
        self.title_prompt = self.load_prompt(self.config['title_prompt_path'])
        
        # 记录失败的文件，用于重试
        self.failed_files_path = os.path.join(self.config['log_path'], 'failed_files.json')
        self.failed_files = self.load_failed_files()
        
        # 打印欢迎信息
        self.print_welcome()

    def load_config(self, path):
        """加载配置文件

        Args:
            path (str): 配置文件路径
            
        Returns:
            dict: 配置信息
        """
        try:

            # with open(path, 'r', encoding='utf-8') as f:
            #     config = json.load(f)
            with resources.open_text(__package__, path) as file:
                config = json.load(file)
                
            # 验证必要的配置项是否存在
            # required_fields = ['api_type', 'api_key', 'model', 'input_folder',
            #                  'output_folder', 'prompt_path', 'title_prompt_path']
            # for field in required_fields:
            #     if field not in config:
            #         logging.error(f"配置文件缺少必要字段: {field}")
            #         sys.exit(1)
                    
            # 设置默认值
            if 'max_tokens' not in config:
                config['max_tokens'] = 4000
            if 'temperature' not in config:
                config['temperature'] = 0.7
            if 'retry_count' not in config:
                config['retry_count'] = 3
            if 'log_path' not in config:
                config['log_path'] = 'logs'
            if 'report_name' not in config:
                config['report_name'] = 'processing_report.md'
            if 'output_formats' not in config:
                config['output_formats'] = ["markdown", "docx"]
                
            return config
            
        except FileNotFoundError:
            print(f"错误: 配置文件 {path} 不存在")
            sys.exit(1)
        except json.JSONDecodeError:
            print(f"错误: 配置文件 {path} 格式不正确")
            sys.exit(1)

    def load_prompt(self, path):
        """加载提示词文件

        Args:
            path (str): 提示词文件路径
            
        Returns:
            str: 提示词内容
        """
        try:
            print(f"当前包名: {__package__}")
            print(f"尝试加载的文件路径: {path}")
            return resources.read_text(__package__, path)
            # with open(path, 'r', encoding='utf-8') as f:
            #     return f.read()
        except FileNotFoundError:
            logging.error(f"提示词文件不存在: {path}")
            sys.exit(1)

    def ensure_directories(self):
        """确保必要的目录结构存在"""
        # 需要确保存在的目录
        dirs = [
            self.config['log_path']
        ]
        
        for directory in dirs:
            os.makedirs(directory, exist_ok=True)
            logging.info(f"确保目录存在: {directory}")

    def setup_logging(self):
        """设置日志记录"""
        # 确保日志目录存在
        log_dir = self.config['log_path']
        os.makedirs(log_dir, exist_ok=True)
        
        log_file = os.path.join(log_dir, f'processor_{self.current_time}.log')
        
        # 重置之前的日志配置
        for handler in logging.root.handlers[:]: 
            logging.root.removeHandler(handler)
            
        # 新的日志配置
        logging.basicConfig(
            level=logging.INFO,
            format='%(asctime)s - %(levelname)s - %(message)s',
            handlers=[
                logging.FileHandler(log_file, encoding='utf-8', mode='w'),
                logging.StreamHandler()
            ]
        )
        
        # 测试日志文件是否可写
        try:
            logging.info("==== 医学讲座转写稿批量处理工具启动 ====")
            logging.info(f"配置文件: {self.config.get('api_type', 'unknown')}, 模型: {self.config.get('model', 'unknown')}")
            logging.info(f"日志文件位置: {log_file}")
        except Exception as e:
            print(f"\n警告: 日志文件写入失败: {str(e)}")

    def load_failed_files(self):
        """加载之前处理失败的文件记录
        
        Returns:
            list: 失败文件列表
        """
        if not os.path.exists(self.failed_files_path):
            return []
            
        try:
            with open(self.failed_files_path, 'r', encoding='utf-8') as f:
                return json.load(f)
        except (json.JSONDecodeError, FileNotFoundError):
            return []

    def save_failed_files(self, failed_list):
        """保存处理失败的文件记录
        
        Args:
            failed_list (list): 失败文件列表
        """
        with open(self.failed_files_path, 'w', encoding='utf-8') as f:
            json.dump(failed_list, f, ensure_ascii=False, indent=2)

    def print_welcome(self):
        """打印欢迎信息"""
        print("\n" + "="*60)
        print("        医学讲座转写稿批量处理工具")
        print("="*60)
        print(f"• 输入文件夹: {self.config['input_folder']}")
        print(f"• 输出文件夹: {self.config['output_folder']}")
        print(f"• 使用的模型: {self.config['model']}")
        print(f"• 处理模式: {'仅重试失败文件' if self.retry_only else '处理所有文件'}")
        print("="*60 + "\n")

    def process_all(self):
        """批量处理所有文件"""
        logging.info("开始批量处理转写稿")
        
        # 初始化报告
        report = {
            'start_time': self.current_time,
            'end_time': '',
            'processed': 0,
            'succeeded': 0,
            'failed': 0,
            'details': []
        }
        
        # 收集要处理的文件
        files_to_process = []
        
        # 根据处理模式确定要处理的文件
        if self.retry_only and self.failed_files:
            logging.info(f"重试模式: 将处理 {len(self.failed_files)} 个之前失败的文件")
            for file in self.failed_files:
                if os.path.exists(os.path.join(self.config['input_folder'], file)):
                    files_to_process.append(file)
        else:
            # 处理所有.txt和.docx文件
            for filename in os.listdir(self.config['input_folder']):
                if filename.endswith('.txt') or filename.endswith('.docx'):
                    files_to_process.append(filename)
            logging.info(f"标准模式: 找到 {len(files_to_process)} 个要处理的文件")
        
        # 新的失败文件列表
        new_failed_files = []
        
        # 处理每个文件
        total_files = len(files_to_process)
        for index, filename in enumerate(files_to_process, 1):
            print(f"\n[{index}/{total_files}] 处理文件: {filename}")
            result = self.process_file(filename)
            report['details'].append(result)
            report['processed'] += 1
            
            if result['status'] == 'success':
                report['succeeded'] += 1
                print(f"✓ 成功处理: {filename}")
            else:
                report['failed'] += 1
                new_failed_files.append(filename)
                print(f"✗ 处理失败: {filename} - {result.get('error', '未知错误')}")
        
        # 更新失败文件列表
        self.save_failed_files(new_failed_files)
        
        # 完成报告
        report['end_time'] = datetime.now().strftime('%Y%m%d_%H%M%S')
        self.generate_report(report)
        
        # 打印总结
        print("\n" + "="*60)
        print(f"处理完成! 总计: {report['processed']}个文件")
        print(f"• 成功: {report['succeeded']}个")
        print(f"• 失败: {report['failed']}个")
        print(f"• 报告已保存: {self.config['report_name']}")
        print("="*60 + "\n")
        
        return report

    def extract_expert_info(self, response):
        """从API响应中提取专家信息
        
        Args:
            response (str): API响应文本
            
        Returns:
            dict: 包含主讲者、主持人和点评专家信息的字典
        """
        # 添加详细日志，记录提取过程
        logging.info("开始提取专家信息...")
        
        # 记录AI响应的前500个字符，方便排查问题
        logging.info(f"AI响应内容前500字符:\n{response[:500]}...")
        
        expert_info = {
            'main_speaker': {
                'name': '',
                'affiliation': '',
                'department': '',
                'title': ''
            },
            'host': {
                'name': '',
                'affiliation': '',
                'department': '',
                'title': ''
            },
            'commentators': []
        }
        
        # 记录原始API响应中有关专家信息的部分
        main_speaker_info = re.search(r'\*\*主讲者信息\*\*[\uff1a:](.*?)(?=\*\*点评专家信息\*\*|$)', response, re.DOTALL)
        if main_speaker_info:
            logging.info(f"找到主讲者信息块:\n{main_speaker_info.group(0)}")
            main_info = main_speaker_info.group(1).strip()
            
            # 提取主讲者各项信息
            name_match = re.search(r'姓名[\uff1a:](.*?)(?=单位[\uff1a:]|科室[\uff1a:]|职称[\uff1a:]|$)', main_info, re.DOTALL)
            if name_match and name_match.group(1).strip():
                expert_info['main_speaker']['name'] = name_match.group(1).strip()
                logging.info(f"提取到主讲者姓名: {expert_info['main_speaker']['name']}")
            
            affiliation_match = re.search(r'单位[\uff1a:](.*?)(?=姓名[\uff1a:]|科室[\uff1a:]|职称[\uff1a:]|$)', main_info, re.DOTALL)
            if affiliation_match and affiliation_match.group(1).strip():
                expert_info['main_speaker']['affiliation'] = affiliation_match.group(1).strip()
                logging.info(f"提取到主讲者单位: {expert_info['main_speaker']['affiliation']}")
            
            department_match = re.search(r'科室[\uff1a:](.*?)(?=姓名[\uff1a:]|单位[\uff1a:]|职称[\uff1a:]|$)', main_info, re.DOTALL)
            if department_match and department_match.group(1).strip():
                expert_info['main_speaker']['department'] = department_match.group(1).strip()
                logging.info(f"提取到主讲者科室: {expert_info['main_speaker']['department']}")
            
            title_match = re.search(r'职称[\uff1a:](.*?)(?=姓名[\uff1a:]|单位[\uff1a:]|科室[\uff1a:]|$)', main_info, re.DOTALL)
            if title_match and title_match.group(1).strip():
                expert_info['main_speaker']['title'] = title_match.group(1).strip()
                logging.info(f"提取到主讲者职称: {expert_info['main_speaker']['title']}")
        else:
            logging.info("未找到主讲者信息块")
        
        # 提取主持人信息
        host_info = re.search(r'\*\*主持人信息\*\*[\uff1a:](.*?)(?=\*\*主讲者信息\*\*|\*\*点评专家信息\*\*|$)', response, re.DOTALL)
        if host_info:
            logging.info(f"找到主持人信息块:\n{host_info.group(0)}")
            host_info_text = host_info.group(1).strip()
            
            # 检查是否是引用其他角色
            if "无专门主持人" in host_info_text or "无" == host_info_text.strip():
                logging.info("主持人信息标注为'无专门主持人'")
            elif "同主讲者" in host_info_text:
                logging.info("主持人信息标注为'同主讲者'")
                expert_info['host'] = expert_info['main_speaker'].copy()
            elif "同点评专家" in host_info_text:
                logging.info("主持人信息标注为'同点评专家'")
                # 注意：此时可能点评专家尚未提取，稍后处理
            else:
                # 提取主持人各项信息
                name_match = re.search(r'姓名[\uff1a:](.*?)(?=单位[\uff1a:]|科室[\uff1a:]|职称[\uff1a:]|$)', host_info_text, re.DOTALL)
                if name_match and name_match.group(1).strip():
                    expert_info['host']['name'] = name_match.group(1).strip()
                    logging.info(f"提取到主持人姓名: {expert_info['host']['name']}")
                
                affiliation_match = re.search(r'单位[\uff1a:](.*?)(?=姓名[\uff1a:]|科室[\uff1a:]|职称[\uff1a:]|$)', host_info_text, re.DOTALL)
                if affiliation_match and affiliation_match.group(1).strip():
                    expert_info['host']['affiliation'] = affiliation_match.group(1).strip()
                    logging.info(f"提取到主持人单位: {expert_info['host']['affiliation']}")
                
                department_match = re.search(r'科室[\uff1a:](.*?)(?=姓名[\uff1a:]|单位[\uff1a:]|职称[\uff1a:]|$)', host_info_text, re.DOTALL)
                if department_match and department_match.group(1).strip():
                    expert_info['host']['department'] = department_match.group(1).strip()
                    logging.info(f"提取到主持人科室: {expert_info['host']['department']}")
                
                title_match = re.search(r'职称[\uff1a:](.*?)(?=姓名[\uff1a:]|单位[\uff1a:]|科室[\uff1a:]|$)', host_info_text, re.DOTALL)
                if title_match and title_match.group(1).strip():
                    expert_info['host']['title'] = title_match.group(1).strip()
                    logging.info(f"提取到主持人职称: {expert_info['host']['title']}")
        else:
            logging.info("未找到主持人信息块")
        
        # 提取点评专家信息
        commentator_pattern = r'\*\*点评专家信息\*\*[\uff1a:](.*?)(?=\n\n|$)'
        commentator_match = re.search(commentator_pattern, response, re.DOTALL)
        if commentator_match:
            logging.info(f"找到点评专家信息块:\n{commentator_match.group(0)}")
            commentator_info = commentator_match.group(1).strip()
            
            # 分析多位专家，先尝试匹配标记的专家
            commentator_blocks = re.findall(r'专家\d+[\uff1a:](.*?)(?=专家\d+[\uff1a:]|$)', commentator_info + '专家999:', re.DOTALL)
            
            # 如果没有找到标记的专家，则假设只有一位未标记的点评专家
            if not commentator_blocks:
                commentator_blocks = [commentator_info]
                logging.info("未找到明确标记的专家，假设只有一位点评专家")
            else:
                logging.info(f"找到 {len(commentator_blocks)} 位点评专家")
            
            for block in commentator_blocks:
                commentator = {
                    'name': '',
                    'affiliation': '',
                    'department': '',
                    'title': ''
                }
                
                # 提取点评专家各项信息
                name_match = re.search(r'姓名[\uff1a:](.*?)(?=单位[\uff1a:]|科室[\uff1a:]|职称[\uff1a:]|$)', block, re.DOTALL)
                if name_match and name_match.group(1).strip():
                    commentator['name'] = name_match.group(1).strip()
                    logging.info(f"提取到点评专家姓名: {commentator['name']}")
                
                affiliation_match = re.search(r'单位[\uff1a:](.*?)(?=姓名[\uff1a:]|科室[\uff1a:]|职称[\uff1a:]|$)', block, re.DOTALL)
                if affiliation_match and affiliation_match.group(1).strip():
                    commentator['affiliation'] = affiliation_match.group(1).strip()
                    logging.info(f"提取到点评专家单位: {commentator['affiliation']}")
                
                department_match = re.search(r'科室[\uff1a:](.*?)(?=姓名[\uff1a:]|单位[\uff1a:]|职称[\uff1a:]|$)', block, re.DOTALL)
                if department_match and department_match.group(1).strip():
                    commentator['department'] = department_match.group(1).strip()
                    logging.info(f"提取到点评专家科室: {commentator['department']}")
                
                title_match = re.search(r'职称[\uff1a:](.*?)(?=姓名[\uff1a:]|单位[\uff1a:]|科室[\uff1a:]|$)', block, re.DOTALL)
                if title_match and title_match.group(1).strip():
                    commentator['title'] = title_match.group(1).strip()
                    logging.info(f"提取到点评专家职称: {commentator['title']}")
                
                # 只添加有信息的专家
                if any(commentator.values()):
                    expert_info['commentators'].append(commentator)
        else:
            logging.info("未找到点评专家信息块")
            
        # 处理主持人「同点评专家」的情况
        if expert_info['host']['name'] == '' and expert_info['commentators'] and "同点评专家" in response:
            expert_info['host'] = expert_info['commentators'][0].copy()
            logging.info("已将第一位点评专家信息复制到主持人信息中")
            
        # 记录专家信息提取结果
        has_expert_info = (
            expert_info['main_speaker']['name'] != '' or 
            expert_info['host']['name'] != '' or 
            expert_info['commentators']
        )
        
        if has_expert_info:
            logging.info(f"成功提取专家信息: {json.dumps(expert_info, ensure_ascii=False)}")
        else:
            logging.info("未提取到有效专家信息")
            
        return expert_info
        
    def parse_titles_from_response(self, response):
        """从API响应中解析出标题
        
        Args:
            response (str): API响应文本
            
        Returns:
            dict: 解析出的标题字典
        """
        logging.info("开始解析标题...")
        
        # 记录原始API响应的前200个字符，方便调试
        logging.info(f"原始API响应前部分: {response[:200]}...")
        
        titles = {
            'academic': [],
            'clinical': [],
            'public': []
        }
        
        # 解析学术版标题
        academic1_match = re.search(r'\[基础学术版-1\](.*?)(?=\[|$)', response, re.DOTALL)
        if academic1_match:
            titles['academic'].append(academic1_match.group(1).strip())
            logging.info(f"找到学术版标题1: {academic1_match.group(1).strip()}")
        else:
            logging.info("未找到学术版标题1，尝试替代格式")
            # 尝试替代格式查找
            alt_match = re.search(r'学术版标题一[\s:：]+(.*?)(?=\n|$)', response)
            if alt_match:
                titles['academic'].append(alt_match.group(1).strip())
                logging.info(f"使用替代格式找到学术版标题1: {alt_match.group(1).strip()}")
        
        academic2_match = re.search(r'\[基础学术版-2\](.*?)(?=\[|$)', response, re.DOTALL)
        if academic2_match:
            titles['academic'].append(academic2_match.group(1).strip())
            logging.info(f"找到学术版标题2: {academic2_match.group(1).strip()}")
        else:
            logging.info("未找到学术版标题2，尝试替代格式")
            # 尝试替代格式查找
            alt_match = re.search(r'学术版标题二[\s:：]+(.*?)(?=\n|$)', response)
            if alt_match:
                titles['academic'].append(alt_match.group(1).strip())
                logging.info(f"使用替代格式找到学术版标题2: {alt_match.group(1).strip()}")
        
        # 解析临床决策版标题
        clinical1_match = re.search(r'\[临床决策版-1\](.*?)(?=\[|$)', response, re.DOTALL)
        if clinical1_match:
            titles['clinical'].append(clinical1_match.group(1).strip())
            logging.info(f"找到临床决策版标题1: {clinical1_match.group(1).strip()}")
        else:
            logging.info("未找到临床决策版标题1，尝试替代格式")
            # 尝试替代格式查找
            alt_match = re.search(r'临床版标题一[\s:：]+(.*?)(?=\n|$)', response)
            if alt_match:
                titles['clinical'].append(alt_match.group(1).strip())
                logging.info(f"使用替代格式找到临床决策版标题1: {alt_match.group(1).strip()}")
        
        clinical2_match = re.search(r'\[临床决策版-2\](.*?)(?=\[|$)', response, re.DOTALL)
        if clinical2_match:
            titles['clinical'].append(clinical2_match.group(1).strip())
            logging.info(f"找到临床决策版标题2: {clinical2_match.group(1).strip()}")
        else:
            logging.info("未找到临床决策版标题2，尝试替代格式")
            # 尝试替代格式查找
            alt_match = re.search(r'临床版标题二[\s:：]+(.*?)(?=\n|$)', response)
            if alt_match:
                titles['clinical'].append(alt_match.group(1).strip())
                logging.info(f"使用替代格式找到临床决策版标题2: {alt_match.group(1).strip()}")
        
        # 解析传播增强版标题
        public1_match = re.search(r'\[传播增强版-1\](.*?)(?=\[|$)', response, re.DOTALL)
        if public1_match:
            titles['public'].append(public1_match.group(1).strip())
            logging.info(f"找到传播增强版标题1: {public1_match.group(1).strip()}")
        else:
            logging.info("未找到传播增强版标题1，尝试替代格式")
            # 尝试替代格式查找
            alt_match = re.search(r'传播版标题一[\s:：]+(.*?)(?=\n|$)', response)
            if alt_match:
                titles['public'].append(alt_match.group(1).strip())
                logging.info(f"使用替代格式找到传播增强版标题1: {alt_match.group(1).strip()}")
        
        public2_match = re.search(r'\[传播增强版-2\](.*?)(?=\[|$)', response, re.DOTALL)
        if public2_match:
            titles['public'].append(public2_match.group(1).strip())
            logging.info(f"找到传播增强版标题2: {public2_match.group(1).strip()}")
        else:
            logging.info("未找到传播增强版标题2，尝试替代格式")
            # 尝试替代格式查找
            alt_match = re.search(r'传播版标题二[\s:：]+(.*?)(?=\n|$)', response)
            if alt_match:
                titles['public'].append(alt_match.group(1).strip())
                logging.info(f"使用替代格式找到传播增强版标题2: {alt_match.group(1).strip()}")
        
        # 确保每种类型有两个标题
        for key in titles:
            while len(titles[key]) < 2:
                default_title = f"{key.capitalize()} Title {len(titles[key])+1}"
                titles[key].append(default_title)
                logging.info(f"使用默认标题补充{key}类型标题: {default_title}")
        
        logging.info(f"标题解析完成: {json.dumps(titles, ensure_ascii=False)}")
        return titles
    
    def parse_editors_note_from_response(self, response):
        """从API响应中解析出编者按
        
        Args:
            response (str): API响应文本
            
        Returns:
            str: 解析出的编者按
        """
        logging.info("开始解析编者按...")
        
        # 记录原始API响应中与编者按相关的部分
        editors_section = re.search(r'\[编者按\](.*?)(?=\[|$)', response, re.DOTALL)
        if editors_section:
            logging.info("找到[编者按]标记格式")
            # 提取内容并删除前后空白
            note = editors_section.group(1).strip()
            logging.info(f"编者按内容(前50字): {note[:50]}...")
            logging.info(f"编者按长度: {len(note)} 字符")
            return note
        else:
            logging.info("未找到[编者按]标记格式, 尝试替代格式")
            
            # 尝试替代格式查找
            alt_formats = [
                r'编者按[\s:\uff1a]+(.*?)(?=\n\n|$)',
                r'编者按\n(.*?)(?=\n\n|$)',
                r'编者按\s*\n\s*([\s\S]*?)(?=\n\n|$)'
            ]
            
            for pattern in alt_formats:
                alt_match = re.search(pattern, response, re.DOTALL)
                if alt_match:
                    note = alt_match.group(1).strip()
                    logging.info(f"使用替代格式找到编者按(前50字): {note[:50]}...")
                    logging.info(f"编者按长度: {len(note)} 字符")
                    return note
            
            # 如果所有格式都未找到
            logging.info("所有尝试的格式都未能找到编者按")
            return ""
        
    def process_file(self, filename):
        """处理单个文件（两步法流程）
        
        Args:
            filename (str): 要处理的文件名
            
        Returns:
            dict: 处理结果
        """
        #file_path = os.path.join(self.config['input_folder'], filename)
        file_path = filename
        result = {
            'filename': filename,
            'start_time': datetime.now().isoformat(),
            'status': 'success'
        }

        try:
            logging.info(f"开始处理文件: {filename}")
            
            # 读取转写稿内容
            if filename.endswith('.txt'):
                with open(file_path, 'r', encoding='utf-8') as f:
                    transcript = f.read()
            elif filename.endswith('.docx'):
                logging.info(f"读取Word文档: {filename}")
                doc = Document(file_path)
                transcript = '\n'.join([paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()])

            # 检查文本长度
            text_length = len(transcript)
            logging.info(f"文本长度: {text_length} 字符")
            
            # 第一步：生成文章主体内容
            logging.info("第一步：生成文章主体内容")
            
            # 根据文本长度决定处理方式
            if text_length > 250000:  # 25万字以上的超长文本
                logging.info("检测到超长文本，将使用分块处理")
                # 分块处理超长文本，并获取专家信息
                article_content, expert_info = self.process_long_text_with_expert_info(transcript)
            else:
                # 常规长度文本处理
                # 读取医学提示词
                medical_prompt = self.load_prompt(self.config['prompt_path'])
                
                # 调用API生成内容
                response = self.generate_content(transcript, medical_prompt, params=True)
                
                # 提取专家信息
                expert_info = self.extract_expert_info(response)
                
                # 清理内容，移除专家信息部分
                article_content = re.sub(r'\*\*主讲者信息\*\*[\uff1a:].*?(?=\n\n|\*\*点评专家信息\*\*|$)', '', 
                                         response, flags=re.DOTALL)
                article_content = re.sub(r'\*\*点评专家信息\*\*[\uff1a:].*?(?=\n\n|$)', '', 
                                         article_content, flags=re.DOTALL)
                article_content = re.sub(r'\n{3,}', '\n\n', article_content).strip()
            
            # 记录专家信息（如有）
            if expert_info['main_speaker']['name'] or expert_info['commentators']:
                logging.info(f"提取到专家信息: {json.dumps(expert_info, ensure_ascii=False)}")
                main_speaker = expert_info['main_speaker']
                logging.info(f"主讲者: {main_speaker.get('name', '')} {main_speaker.get('affiliation', '')} "
                             f"{main_speaker.get('department', '')} {main_speaker.get('title', '')}")
                
                if expert_info['commentators']:
                    for i, commentator in enumerate(expert_info['commentators'], 1):
                        logging.info(f"点评专家1: {commentator.get('name', '')} {commentator.get('affiliation', '')} "
                                     f"{commentator.get('department', '')} {commentator.get('title', '')}")
            else:
                logging.info("未提取到专家信息")
            
            # 第二步：生成标题和编者按
            logging.info("第二步：生成标题和编者按")
            titles, editors_note = self.generate_titles_and_note(article_content, expert_info)
            
            # 使用clean_content_after_editors_note函数重构文章内容
            logging.info("调用clean_content_after_editors_note函数重构文章内容...")
            restructured_content = self.clean_content_after_editors_note(article_content)
            logging.info(f"重构前内容长度: {len(article_content)}字符, 重构后内容长度: {len(restructured_content)}字符")
            
            # 保存结果
            base_name = os.path.splitext(filename)[0]
            # 处理文件名中的特殊字符
            base_name = base_name.replace('_', ' ').strip()
            # 增加_article后缀以区分原文件
            output_name = f"{base_name}_article"
            #output_path = os.path.join(self.config['output_folder'], output_name)
            output_path = output_name
            
            if "markdown" in self.config['output_formats']:
                self.save_markdown(output_path, restructured_content, titles, editors_note)
            
            if "docx" in self.config['output_formats']:
                result_path = self.save_word(output_path, restructured_content, titles, editors_note)
                result['filename'] = result_path
            logging.info(f"文件处理成功: {filename}")

        except Exception as e:
            result['status'] = 'failed'
            result['error'] = str(e)
            logging.error(f"处理失败: {filename} - {str(e)}")
            # 展示详细错误信息便于调试
            logging.error(traceback.format_exc())

        result['end_time'] = datetime.now().isoformat()
        return result

    def generate_titles_and_note(self, content, expert_info):
        """基于内容和专家信息生成标题和编者按
        
        Args:
            content (str): 文章主体内容
            expert_info (dict): 专家信息字典
            
        Returns:
            tuple: (titles, editors_note) 生成的标题字典和编者按
        """
        logging.info("第二步：基于内容和专家信息生成标题和编者按")
        
        # 记录专家信息详情，以便分析
        logging.info(f"传递给标题生成的专家信息详情: {json.dumps(expert_info, ensure_ascii=False)}")
        
        # 读取标题编者按提示词模板
        title_prompt_path = self.config.get('title_prompt_path', 'title_note_prompt.txt')
        try:
            with resources.open_text("medical_article_sdk", title_prompt_path, encoding='utf-8') as f:
                prompt = f.read()
        except FileNotFoundError:
            logging.error(f"标题编者按提示词文件不存在: {title_prompt_path}")
            raise
            
        # 替换专家信息
        main_speaker = expert_info['main_speaker']
        prompt = prompt.replace("{main_speaker_name}", main_speaker.get('name', ''))
        prompt = prompt.replace("{main_speaker_affiliation}", main_speaker.get('affiliation', ''))
        prompt = prompt.replace("{main_speaker_department}", main_speaker.get('department', ''))
        prompt = prompt.replace("{main_speaker_title}", main_speaker.get('title', ''))
        
        # 处理点评专家信息
        commentators_text = ""
        for i, commentator in enumerate(expert_info.get('commentators', []), 1):
            commentators_text += f"专家{i}：\n"
            commentators_text += f"姓名：{commentator.get('name', '')}\n"
            commentators_text += f"单位：{commentator.get('affiliation', '')}\n"
            commentators_text += f"科室：{commentator.get('department', '')}\n"
            commentators_text += f"职称：{commentator.get('title', '')}\n\n"
        
        if not commentators_text:
            commentators_text = "无点评专家"
        
        prompt = prompt.replace("{commentators_info}", commentators_text)
        
        # 替换内容
        prompt = prompt.replace("[CONTENT]", content)
        
        # 记录提示词大小
        logging.info(f"标题编者按提示词大小: {len(prompt)} 字符")
        logging.info(f"标题编者按提示词前100字符: {prompt[:100]}...")
        
        # 调用API生成内容，明确指定使用标题和编者按生成参数
        response = self.generate_content(content[:3000], prompt, params=False)  # 只使用内容前部分，避免过长
        
        # 记录完整响应以便分析
        logging.info(f"标题和编者按生成API完整响应:\n{response}\n{'='*80}")
        
        # 解析标题和编者按
        titles = self.parse_titles_from_response(response)
        editors_note = self.parse_editors_note_from_response(response)
        
        return titles, editors_note
        
    def generate_content(self, text, prompt, params=None, max_retries=3, timeout=120):
        """调用API生成内容
        
        Args:
            text (str): 原始文本
            prompt (str): 提示词
            params (dict|bool): 生成参数，可以是字典或布尔值（兼容旧版调用）
                               如果是布尔值，则True表示内容生成，False表示标题生成
            max_retries (int): 最大重试次数
            timeout (int): 请求超时时间（秒）
            
        Returns:
            str: 生成的内容
            
        Raises:
            Exception: 当所有重试都失败时抛出异常
        """
        # 兼容两种不同的调用方式
        if params is None:
            # 默认使用内容生成参数
            temperature = self.config.get('content_temp', 0.3)
            logging.info(f"使用默认内容生成参数: 温度={temperature}")
        elif isinstance(params, bool):
            # 如果是布尔值（兼容旧版调用）
            if params:  # params 为 True，表示内容生成
                temperature = self.config.get('content_temp', 0.3)
                logging.info(f"使用内容生成参数: 温度={temperature}")
            else:  # params 为 False，表示标题编者按生成
                temperature = self.config.get('title_note_temp', 0.65)
                logging.info(f"使用标题编者按生成参数: 温度={temperature}")
        elif isinstance(params, dict):
            # 如果是字典，直接使用其中的temperature值
            temperature = params.get('temperature', 0.7)
            logging.info(f"使用指定生成参数: 温度={temperature}")
        else:
            # 非法参数类型，使用默认值
            temperature = 0.7
            logging.info(f"使用默认参数: 温度={temperature}")
        
        # 组合最终提示词
        final_prompt = f"{prompt}\n\n原文内容:\n{text}"
        
        # 记录终极提示词长度
        prompt_length = len(final_prompt)
        logging.info(f"提示词长度: {prompt_length} 字符")
        
        # 设置API请求参数
        model = self.config.get('model_name', 'doubao-1-5-pro-256k-250115')  # 默认值
        
        # 重试机制
        for attempt in range(1, max_retries + 1):
            logging.info(f"API调用尝试 {attempt}/{max_retries}")
            try:
                # 根据配置选择API类型
                if self.config.get('api_type', 'volcengine') == 'volcengine':
                    response = self._call_volcengine_api(final_prompt, model, temperature, timeout)
                else:
                    raise ValueError(f"不支持的API类型: {self.config.get('api_type')}")
                
                return response
                
            except Exception as e:
                if 'timeout' in str(e).lower():
                    logging.warning(f"API请求错误: {str(e)}")
                    if attempt < max_retries:
                        retry_delay = 3  # 重试间隔
                        logging.info(f"等待 {retry_delay} 秒后重试...")
                        time.sleep(retry_delay)
                    else:
                        logging.error(f"所有重试得到相同错误: {str(e)}")
                        raise
                else:
                    logging.error(f"API调用错误: {str(e)}")
                    raise
        
        raise Exception(f"达到最大重试次数 {max_retries}")
    
    def _call_volcengine_api(self, prompt, model, temperature, timeout):
        """调用火山引擎API生成内容
        
        Args:
            prompt (str): 完整提示词
            model (str): 模型名称
            temperature (float): 生成温度
            timeout (int): 超时时间（秒）
            
        Returns:
            str: API响应生成的内容
        """
        # API请求URL和头信息
        url = "https://ark.cn-beijing.volces.com/api/v3/chat/completions"
        headers = {
            "Content-Type": "application/json",
            "Authorization": f"Bearer {self.api_key}"
        }
        
        # 请求体
        payload = {
            "model": model,
            "messages": [
                {
                    "role": "system",
                    "content": "你是专业的医学内容生成助手。"
                },
                {
                    "role": "user",
                    "content": prompt
                }
            ],
            "temperature": temperature,
            "top_p": 0.95
        }
        
        # 发送请求
        response = requests.post(url, headers=headers, json=payload, timeout=timeout)
        
        # 检查响应状态
        if response.status_code != 200:
            error_message = f"火山引擎API错误: 状态码 {response.status_code}, 响应: {response.text}"
            logging.error(error_message)
            raise Exception(error_message)
        
        # 解析响应
        result = response.json()
        
        # 检查是否成功
        logging.debug(f"API响应：{json.dumps(result, ensure_ascii=False)[:200]}...")
        
        # 提取生成的内容
        if 'choices' not in result:
            error_message = f"火山引擎API响应格式错误: {json.dumps(result, ensure_ascii=False)}"
            logging.error(error_message)
            raise Exception(error_message)
            
        generated_content = result.get('choices', [{}])[0].get('message', {}).get('content', '')
        
        # 记录生成的内容长度和前100个字符
        content_length = len(generated_content)
        logging.info(f"生成的内容长度: {content_length} 字符")
        logging.info(f"生成的内容开头(前100字符): {generated_content[:100]}...")
        
        return generated_content
        




    def process_long_text_with_expert_info(self, text, max_chunk_size=100000):
        """处理超长文本并提取专家信息

        将超长文本分块处理，然后合并结果，并提取专家信息
        
        Args:
            text (str): 原始文本
            max_chunk_size (int): 每块最大字符数
            
        Returns:
            tuple: (content, expert_info) 处理后的文章内容和专家信息
        """
        logging.info(f"将长文本分为多个块处理，每块最大 {max_chunk_size} 字符")
        
        # 分块
        chunks = [text[i:i+max_chunk_size] for i in range(0, len(text), max_chunk_size)]
        chunk_results = []
        expert_info = {
            'main_speaker': {
                'name': '',
                'affiliation': '',
                'department': '',
                'title': ''
            },
            'commentators': []
        }
        
        # 处理每个块
        for i, chunk in enumerate(chunks):
            logging.info(f"处理第 {i+1}/{len(chunks)} 块")
            # 为每个块添加特殊指令
            medical_prompt = self.load_prompt(self.config['prompt_path'])
            chunk_prompt = f"{medical_prompt}\n\n注意：这是一个分块处理的片段，是第 {i+1}/{len(chunks)} 块。"
            response = self.generate_content(chunk, chunk_prompt, params=True)
            
            # 从第一块提取专家信息
            if i == 0:
                expert_info = self.extract_expert_info(response)
                logging.info(f"从第一块提取到的专家信息: {json.dumps(expert_info, ensure_ascii=False)}")
            
            # 清理内容，移除专家信息部分
            clean_content = re.sub(r'\*\*主讲者信息\*\*[\uff1a:].*?(?=\n\n|\*\*点评专家信息\*\*|$)', '', 
                                   response, flags=re.DOTALL)
            clean_content = re.sub(r'\*\*点评专家信息\*\*[\uff1a:].*?(?=\n\n|$)', '', 
                                   clean_content, flags=re.DOTALL)
            chunk_results.append(clean_content.strip())
        
        # 合并处理结果
        if len(chunk_results) == 1:
            return chunk_results[0], expert_info
        
        # 构建合并提示词
        merge_prompt = """你是一位精通医学的内容编辑专家。我将提供一个已经被分成多个部分处理的医学讲座文章。
请将这些部分整合成一篇连贯的医学文章，符合以下要求：
1. 删除各部分之间的重复内容
2. 确保论点和主题的连贯性
3. 维持正确的标题层级和逻辑结构
4. 最终文章应当是一篇完整的、符合医学格式的专业文章
5. 文章总长度应保持在2500-3000字左右
6. 保留所有重要的临床知识点和实践指导

以下是需要整合的文章部分："""
        
        # 合并所有部分
        merged_text = "\n\n===第1部分===\n\n" + \
                   "\n\n===第2部分===\n\n".join(chunk_results)
        
        merged_result = self.generate_content(merged_text, merge_prompt, params=True)
        
        logging.info("长文本分块处理完成并合并")
        return merged_result, expert_info
        
    def generate_titles(self, article):
        """生成三种风格的标题，每种两个
        
        Args:
{{ ... }}
            article (str): 文章内容
            
        Returns:
            dict: 三种风格的标题，每种两个
        """
        
        # 标题生成提示词
        title_types = {
            'academic': {
                'prompt': f"{self.title_prompt}\n\n请生成两个学术风格的标题，适合在医学期刊中发表，应当准确反映内容、简洁专业。\n\n请将每个标题分行返回，格式为：\n学术版标题一：[标题内容]\n学术版标题二：[标题内容]",
                'params': {'temperature': 0.5}  # 学术标题需要更严谨，使用较低的temperature
            },
            'clinical': {
                'prompt': f"{self.title_prompt}\n\n请生成两个临床决策版标题，面向临床医生，强调实用性和临床决策价值。\n\n请将每个标题分行返回，格式为：\n临床版标题一：[标题内容]\n临床版标题二：[标题内容]",
                'params': {'temperature': 0.7}  # 平衡创造性和实用性
            },
            'public': {
                'prompt': f"{self.title_prompt}\n\n请生成两个传播增强版标题，在保持专业准确的同时，更具吸引力和传播性，适合在医学教育平台推广。\n\n请将每个标题分行返回，格式为：\n传播版标题一：[标题内容]\n传播版标题二：[标题内容]",
                'params': {'temperature': 0.9}  # 传播版需要更有创意，使用较高的temperature
            }
        }
        
        titles = {}
        
        # 生成每种类型的标题
        for title_type, config in title_types.items():
            logging.info(f"生成{title_type}标题...")
            title_response = self.generate_content(
                article[:3000],  # 只使用文章前部分生成标题
                config['prompt'],
                config['params']
            )
            
            # 分析返回的两个标题
            title_lines = [line.strip() for line in title_response.split('\n') if line.strip()]
            
            # 初始化标题列表
            titles[title_type] = []
            
            # 解析每个标题
            for line in title_lines:
                if ':' in line or '：' in line:  # 处理中文冒号和英文冒号
                    # 提取实际标题内容（删除前缀）
                    if ':' in line:
                        title_content = line.split(':', 1)[1].strip().strip('"').strip("'")
                    else:
                        title_content = line.split('：', 1)[1].strip().strip('"').strip("'")
                    
                    titles[title_type].append(title_content)
                else:
                    # 如果没有按要求格式返回，将整行作为标题
                    titles[title_type].append(line.strip().strip('"').strip("'"))
            
            # 确保有两个标题（如果AI返回格式不正确）
            while len(titles[title_type]) < 2:
                titles[title_type].append(f"{title_type.capitalize()} Title {len(titles[title_type])+1}")
            
            # 限制为2个标题
            titles[title_type] = titles[title_type][:2]
            
            # 记录到日志
            logging.info(f"{title_type}标题一: {titles[title_type][0]}")
            logging.info(f"{title_type}标题二: {titles[title_type][1]}")
            
        return titles

    def save_markdown(self, path, content, titles, editors_note):
        """保存为Markdown格式
        
        Args:
            path (str): 输出路径（不含扩展名）
            content (str): 文章内容
            titles (dict): 标题字典
            editors_note (str): 编者按
        """
        import re
        logging.info(f"保存Markdown文件: {path}.md")
        
        # 第二层防护：再次扫描内容，移除多余的一级标题
        main_title = titles['academic'][0]  # 主标题
        lines = content.split('\n')
        filtered_lines = []
        found_first_h1 = False
        
        for line in lines:
            # 检测一级标题
            h1_match = re.match(r'^# (.+)$', line.strip())
            if h1_match:
                h1_text = h1_match.group(1).strip()
                if not found_first_h1:
                    found_first_h1 = True
                    filtered_lines.append(line)
                else:
                    # 记录并跳过多余的一级标题
                    logging.info(f"在save_markdown中移除多余的一级标题: {h1_text}")
                    continue
            else:
                filtered_lines.append(line)
        
        # 构建 Markdown 内容
        md_content = f"# {titles['academic'][0]}\n\n"
        
        # 添加编者按（如果有）
        if editors_note and editors_note.strip():
            md_content += f"## 编者按\n{editors_note}\n\n"
            
        md_content += f"## 临床决策版标题\n{titles['clinical'][0]}\n{titles['clinical'][1]}\n\n"
        md_content += f"## 传播增强版标题\n{titles['public'][0]}\n{titles['public'][1]}\n\n"
        md_content += '\n'.join(filtered_lines)
        
        # 确保目录存在
        os.makedirs(os.path.dirname(path), exist_ok=True)
        
        # 写入文件
        with open(f"{path}.md", 'w', encoding='utf-8') as f:
            f.write(md_content)

    def clean_content_after_editors_note(self, content):
        """主动重构文章内容，只保留从第一个二级标题开始的内容
        
        Args:
            content (str): 原始内容
            
        Returns:
            str: 清理后的内容
        """
        logging.info("开始重构文章内容，只保留从第一个二级标题开始的内容...")
        
        lines = content.split('\n')
        h2_content = []
        
        # 查找第一个二级标题及其后的所有内容
        first_h2_index = -1
        for i, line in enumerate(lines):
            if line.strip().startswith('## '):
                first_h2_index = i
                logging.info(f"找到第一个二级标题: {line.strip()}")
                break
        
        if first_h2_index != -1:
            h2_content = lines[first_h2_index:]
            logging.info(f"提取到第一个二级标题及后续内容，共{len(h2_content)}行")
            # 合并所有行成最终内容
            final_content = '\n'.join(h2_content)
            logging.info(f"重构后的内容长度: {len(final_content)}字符")
            return final_content
        else:
            # 如果没找到二级标题，返回空字符串
            logging.warning("未找到任何二级标题，返回空内容")
            return ""
        
    def fix_content_format(self, content):
        """修复内容格式问题，包括标题层级和列表编号
        
        Args:
            content (str): 原始内容
            
        Returns:
            str: 修复过格式问题的内容
        """
        import re
        
        # 按行分割内容
        lines = content.split('\n')
        fixed_lines = []
        
        found_first_h1 = False  # 跟踪是否找到第一个一级标题
        current_section = "default"  # 当前所在的章节名称
        list_counters = {"default": 0}  # 列表计数器，每个章节单独一个
        in_list = False  # 跟踪是否正在处理列表
        
        # 用于匹配列表项的正则表达式（数字+点或括号+空格）
        list_item_regex = re.compile(r'^(\s*)([0-9]+)[.)](\s)')
        
        for i, line in enumerate(lines):
            # 处理空行
            if not line.strip():
                fixed_lines.append(line)
                continue
                
            # 检测标准Markdown标题行
            if line.strip().startswith('#'):
                heading_match = re.match(r'^(#+)\s+(.+)$', line.strip())
                if heading_match:
                    level = len(heading_match.group(1))  # #的数量表示标题级别
                    text = heading_match.group(2).strip()
                    
                    # 处理一级标题
                    if level == 1:
                        # 如果已经找到第一个一级标题，则跳过当前行
                        if found_first_h1:
                            # 记录移除的标题信息
                            logging.info(f"第一层防护：移除多余的一级标题: {text}")
                            # 如果当前行是分隔线之后的标题，更可能是多余的
                            if i > 0:
                                # 检查前面的行是否包含分隔线
                                # 注意：分隔线可能是'---'或者'____________'等格式
                                prev_line = lines[i-1].strip()
                                if prev_line.startswith('---') or all(c == '_' for c in prev_line) or all(c == '-' for c in prev_line):
                                    logging.info(f"第一层防护：移除分隔线之后的标题: {text}")
                            continue  # 不将这一行加入到结果中
                        else:
                            found_first_h1 = True
                            # 保存第一个一级标题的文本，用于后续相似度比较
                            first_h1_text = text
                    
                    # 如果是章节标题（二级或三级），更新当前章节并重置列表计数器
                    if level in [2, 3]:
                        current_section = text
                        list_counters[current_section] = 0  # 新章节开始，重置计数器
                        in_list = False  # 重置列表状态
                    
                    fixed_lines.append(line)  # 添加原始标题行
                else:
                    fixed_lines.append(line)  # 非标准标题格式，原样添加
                
                continue  # 处理完标题行后继续下一行
            
            # 处理列表项
            list_match = list_item_regex.match(line)
            if list_match:
                indent = list_match.group(1)  # 缩进
                spaces = list_match.group(3)  # 空格
                
                # 确保我们有列表计数器给当前章节
                if current_section not in list_counters:
                    list_counters[current_section] = 0
                
                # 增加计数器并重建列表项，确保编号从1开始
                list_counters[current_section] += 1
                rest_of_line = line[list_match.end():]
                fixed_line = f"{indent}{list_counters[current_section]}){spaces}{rest_of_line}"
                
                fixed_lines.append(fixed_line)
                in_list = True
                continue
            
            # 如果是普通文本行且非缩进，可能标志列表结束
            if line.strip() and not line.startswith(' ') and in_list:
                in_list = False
            
            # 其他所有行直接添加
            fixed_lines.append(line)
        
        # 将处理后的行重新组合成字符串
        return '\n'.join(fixed_lines)

    def save_word(self, path, content, titles, editors_note):
        """保存为Word格式
        
        Args:
            path (str): 输出路径（不含扩展名）
            content (str): 文章内容
            titles (dict): 标题字典，每种类型包含两个标题
            editors_note (str): 编者按
        """
        logging.info(f"保存Word文件: {path}.docx")
        
        # 先清理编者按和第一个二级标题之间的无关内容
        clean_content = self.clean_content_after_editors_note(content)
        
        # 再修复内容格式问题
        fixed_content = self.fix_content_format(clean_content)
        
        # 创建Word文档
        doc = Document()
        
        # 记录主标题，作为判断重复一级标题的依据
        main_title = titles['academic'][0]
        
        # 初始化标记变量，表示已添加过一级标题
        self._first_h1_added = True
        
        # 设置文档主标题（使用第一个学术标题）
        logging.info(f"添加学术标题作为文档主标题: {main_title}")
        doc.add_heading(main_title, level=1)
        
        # 添加标题版本部分
        doc.add_heading("文章标题备选方案", level=2)
        
        # 添加学术版标题
        p = doc.add_paragraph()
        p.add_run("学术版标题一：").bold = True
        p.add_run(titles['academic'][0])
        
        p = doc.add_paragraph()
        p.add_run("学术版标题二：").bold = True
        p.add_run(titles['academic'][1])
        
        # 添加临床决策版标题
        p = doc.add_paragraph()
        p.add_run("临床决策版标题一：").bold = True
        p.add_run(titles['clinical'][0])
        
        p = doc.add_paragraph()
        p.add_run("临床决策版标题二：").bold = True
        p.add_run(titles['clinical'][1])
        
        # 添加传播增强版标题
        p = doc.add_paragraph()
        p.add_run("传播增强版标题一：").bold = True
        p.add_run(titles['public'][0])
        
        p = doc.add_paragraph()
        p.add_run("传播增强版标题二：").bold = True
        p.add_run(titles['public'][1])
        
        # 添加编者按（如果有）
        if editors_note and editors_note.strip():
            doc.add_heading("编者按", level=2)
            doc.add_paragraph(editors_note)
        
        # 添加分隔线
        doc.add_paragraph("—" * 30)

        # 处理内容，保留Markdown格式
        # 拆分段落处理
        paragraphs = [p for p in fixed_content.split('\n') if p.strip()]
        i = 0
        in_list = False  # 跟踪列表状态
        list_type = None  # 列表类型：有序或无序
        list_level = 0   # 列表嵌套级别
        current_list = None  # 当前列表对象
        
        while i < len(paragraphs):
            para = paragraphs[i].strip()
            
            # 标题处理
            if para.startswith('# '):
                # 第三层防护：处理一级标题
                heading_text = para[2:].strip()
                
                # 全局静态变量，用于跟踪是否已添加过一级标题
                # 如果save_word方法里还没有这个变量，代表是第一次运行
                if not hasattr(self, '_first_h1_added'):
                    # 首次调用时初始化为False
                    self._first_h1_added = False
                
                # 如果已经添加过一级标题，跳过这个
                if self._first_h1_added:
                    logging.info(f"第三层防护：已有一级标题，跳过当前标题: {heading_text}")
                    i += 1
                    continue
                
                # 检查是否是分隔线后的标题
                is_after_separator = False
                if i > 0:
                    # 往前检查分隔线，最多检查前三行
                    for j in range(1, min(i+1, 4)):
                        if i-j >= 0:
                            prev_para = paragraphs[i-j].strip()
                            # 检查各种可能的分隔线格式
                            if (prev_para.startswith('---') or 
                                all(c == '_' for c in prev_para) or 
                                all(c == '-' for c in prev_para) or
                                prev_para == "*" * len(prev_para) or
                                prev_para == "=" * len(prev_para) or
                                prev_para.count('-') > 10):
                                is_after_separator = True
                                logging.info(f"第三层防护：该标题在分隔线之后: {heading_text}")
                                break
                
                # 如果不是第一个段落或者在分隔线之后，跳过
                if i > 0 or is_after_separator:
                    logging.info(f"第三层防护：跳过非第一个或分隔线后的一级标题: {heading_text}")
                    i += 1
                    continue
                
                # 添加第一个一级标题并标记为已添加
                logging.info(f"第三层防护：添加第一个一级标题到Word文档: {heading_text}")
                doc.add_heading(heading_text, level=1)
                self._first_h1_added = True  # 标记为已添加过一级标题
                in_list = False
            elif para.startswith('## '):
                doc.add_heading(para[3:], level=2)
                in_list = False
            elif para.startswith('### '):
                doc.add_heading(para[4:], level=3)
                in_list = False
            elif para.startswith('#### '):
                doc.add_heading(para[5:], level=4)
                in_list = False
            
            # 列表处理
            elif para.startswith('- ') or para.startswith('* '): 
                list_text = para[2:].strip()
                if not in_list or list_type != 'unordered':
                    in_list = True
                    list_type = 'unordered'
                    current_list = doc.add_paragraph()
                    # 处理列表项中的格式标记
                    parts = self._parse_markdown_formatting(list_text)
                    for format_type, text in parts:
                        run = current_list.add_run(text)
                        if format_type == 'bold':
                            run.bold = True
                        elif format_type == 'italic':
                            run.italic = True
                        elif format_type == 'bold_italic':
                            run.bold = True
                            run.italic = True
                    # 设置项目符号格式
                    current_list.style = 'List Bullet'
                else:
                    # 继续当前列表
                    new_item = doc.add_paragraph()
                    new_item.style = 'List Bullet'
                    # 处理列表项中的格式标记
                    parts = self._parse_markdown_formatting(list_text)
                    for format_type, text in parts:
                        run = new_item.add_run(text)
                        if format_type == 'bold':
                            run.bold = True
                        elif format_type == 'italic':
                            run.italic = True
                        elif format_type == 'bold_italic':
                            run.bold = True
                            run.italic = True
            
            elif re.match(r'^\d+\.\s', para):  # 匹配数字、点、空格开头的序号
                list_text = re.sub(r'^\d+\.\s', '', para)
                if not in_list or list_type != 'ordered':
                    in_list = True
                    list_type = 'ordered'
                    current_list = doc.add_paragraph()
                    # 处理列表项中的格式标记
                    parts = self._parse_markdown_formatting(list_text)
                    for format_type, text in parts:
                        run = current_list.add_run(text)
                        if format_type == 'bold':
                            run.bold = True
                        elif format_type == 'italic':
                            run.italic = True
                        elif format_type == 'bold_italic':
                            run.bold = True
                            run.italic = True
                    # 设置有序列表格式
                    current_list.style = 'List Number'
                else:
                    # 继续当前列表
                    new_item = doc.add_paragraph()
                    new_item.style = 'List Number'
                    # 处理列表项中的格式标记
                    parts = self._parse_markdown_formatting(list_text)
                    for format_type, text in parts:
                        run = new_item.add_run(text)
                        if format_type == 'bold':
                            run.bold = True
                        elif format_type == 'italic':
                            run.italic = True
                        elif format_type == 'bold_italic':
                            run.bold = True
                            run.italic = True
            
            # 普通段落处理
            else:
                # 处理特殊格式（粗体、斜体）
                p = doc.add_paragraph()
                
                # 分析文本中的Markdown格式
                parts = self._parse_markdown_formatting(para)
                
                for format_type, text in parts:
                    run = p.add_run(text)
                    
                    if format_type == 'bold':
                        run.bold = True
                    elif format_type == 'italic':
                        run.italic = True
                    elif format_type == 'bold_italic':
                        run.bold = True
                        run.italic = True
                
                in_list = False
            
            i += 1
        
        # 确保目录存在
        #os.makedirs(os.path.dirname(path), exist_ok=True)
        
        # 保存文档
        result_path = "{path}.docx"
        doc.save(result_path)
        return result_path
    
    def _parse_markdown_formatting(self, text):
        """解析Markdown格式文本中的格式标记
        
        Args:
            text (str): 原始Markdown文本
            
        Returns:
            list: 包含格式类型和文本的列表，如[('normal', 'text'), ('bold', 'bold text')]
        """
        parts = []
        i = 0
        
        while i < len(text):
            # 匹配粗体斜体格式 (***bold italic***)
            if i + 5 < len(text) and text[i:i+3] == '***' and '***' in text[i+3:]:
                start = i + 3
                end = text.find('***', start)
                parts.append(('bold_italic', text[start:end]))
                i = end + 3
            # 匹配粗体格式 (**bold**)
            elif i + 4 < len(text) and text[i:i+2] == '**' and '**' in text[i+2:]:
                start = i + 2
                end = text.find('**', start)
                parts.append(('bold', text[start:end]))
                i = end + 2
            # 匹配斜体格式 (*italic*)
            elif i + 2 < len(text) and text[i] == '*' and '*' in text[i+1:]:
                start = i + 1
                end = text.find('*', start)
                parts.append(('italic', text[start:end]))
                i = end + 1
            else:
                # 收集普通文本，直到下一个格式标记
                start = i
                special_markers = ['***', '**', '*']
                next_marker_pos = float('inf')
                
                for marker in special_markers:
                    pos = text.find(marker, start)
                    if pos != -1 and pos < next_marker_pos:
                        next_marker_pos = pos
                
                if next_marker_pos == float('inf'):
                    parts.append(('normal', text[start:]))
                    break
                else:
                    if start != next_marker_pos:  # 确保有文本要添加
                        parts.append(('normal', text[start:next_marker_pos]))
                    i = next_marker_pos
                    continue
                
        # 如果解析失败，返回原始文本
        if not parts:
            parts.append(('normal', text))
            
        return parts

    def generate_report(self, report):
        """生成处理报告
        
        Args:
            report (dict): 报告数据
        """
        report_path = self.config['report_name']
        logging.info(f"生成处理报告: {report_path}")
        
        # 计算处理时间
        start_time = datetime.strptime(report['start_time'], '%Y%m%d_%H%M%S')
        end_time = datetime.strptime(report['end_time'], '%Y%m%d_%H%M%S')
        duration = end_time - start_time
        
        # 构建报告内容
        report_content = "# 医学讲座转写稿处理报告\n\n"
        report_content += f"## 执行摘要\n\n"
        report_content += f"- **开始时间**: {start_time.strftime('%Y-%m-%d %H:%M:%S')}\n"
        report_content += f"- **结束时间**: {end_time.strftime('%Y-%m-%d %H:%M:%S')}\n"
        report_content += f"- **处理时长**: {duration}\n"
        report_content += f"- **总处理文件**: {report['processed']} 个\n"
        report_content += f"- **成功转换**: {report['succeeded']} 个\n"
        report_content += f"- **失败文件**: {report['failed']} 个\n"
        report_content += f"- **成功率**: {(report['succeeded'] / report['processed'] * 100) if report['processed'] > 0 else 0:.2f}%\n\n"

        # 添加成功文件列表
        report_content += "## 成功处理的文件\n\n"
        success_files = [detail for detail in report['details'] if detail['status'] == 'success']
        if success_files:
            for detail in success_files:
                report_content += f"- **{detail['filename']}**\n"
                report_content += f"  - 处理用时: {(datetime.fromisoformat(detail['end_time']) - datetime.fromisoformat(detail['start_time'])).total_seconds():.2f} 秒\n"
        else:
            report_content += "没有成功处理的文件\n"

        # 添加失败文件列表
        report_content += "\n## 处理失败的文件\n\n"
        failed_files = [detail for detail in report['details'] if detail['status'] == 'failed']
        if failed_files:
            for detail in failed_files:
                report_content += f"- **{detail['filename']}**\n"
                report_content += f"  - 错误信息: {detail.get('error', '未知错误')}\n"
        else:
            report_content += "没有处理失败的文件\n"

        # 保存报告
        with open(report_path, 'w', encoding='utf-8') as f:
            f.write(report_content)


def parse_args():
    parser = argparse.ArgumentParser(description="医学讲座转写稿批量处理工具")
    parser.add_argument("--config", help="配置文件路径 (默认: config.json)", default="config.json")
    parser.add_argument("--retry", help="只处理之前失败的文件", action="store_true")
    parser.add_argument("--file", help="指定单个文件进行处理，而不是整个文件夹")
    return parser.parse_args()


def main():
    """主函数，程序入口点"""
    # 解析命令行参数
    args = parse_args()
    
    try:
        # 创建处理器实例
        processor = MedicalArticleProcessor(
            config_path=args.config,
            retry_only=args.retry
        )
        
        # 处理单个文件或批量处理
        if args.file:
            # 处理单个指定文件
            file_path = os.path.join(processor.config['input_folder'], args.file)
            if not os.path.exists(file_path):
                print(f"\n错误: 指定的文件 '{args.file}' 不存在")
                sys.exit(1)
            processor.process_file(args.file)
            print(f"\n文件 '{args.file}' 处理完成!")
        else:
            # 批量处理所有文件
            processor.process_all()
        
    except Exception as e:
        logging.error(f"程序执行出错: {str(e)}")
        print(f"\n错误: {str(e)}")
        sys.exit(1)


if __name__ == "__main__":
    main()
