# -*- coding: utf-8 -*-

""" Copied and adapted from https://gist.github.com/thiagodiniz/

https://gist.github.com/thiagodiniz/fe5c1417a93bc640358a3e2322b2fc5f
"""

import io

import requests
from bs4 import NavigableString, BeautifulSoup
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.shared import Mm
from docx.image.exceptions import UnrecognizedImageError
from lxml import etree

A4_PAGE_HEIGHT = 297
A4_PAGE_WIDHT = 210
DEFAULT_MARGINS = 25.4
MAX_PICTURE_SIZE = 602


def parse_html_to_docx(html, document):
    soup = BeautifulSoup(html, features="lxml")

    body = soup.body
    remove_newlines(body)
    process_node(body, document)

    return body


def process_node(node, document, parent=None, paragraph=None, style=[]):
    try:
        # if node.contents:
        #     pass

        if node.name in ["ul", "ol"]:
            process_list(node, document)
            return None

        elif node.name == "math":
            process_math(node, document, paragraph)
            return None

        elif node.name == "table":
            process_table(document, node)
            return None

        elif node.name in ["p", "h1", "h2", "h3", "h4", "h5", "h6"]:
            paragraph = document.add_paragraph("")
            style = []

        elif node.name == "img":
            process_image(node, document, paragraph, style)
            return None

        elif node.name == "br":
            document.add_paragraph("")
            return None

        if node.get("style") is not None:
            style = style + process_style(node["style"])

        for child in node.children:
            paragraph = process_node(child, document, node, paragraph, style)

    except AttributeError:
        process_string(node, document, parent, paragraph, style)

    return paragraph


def process_string(str_node, document, parent, paragraph, style=[]):
    if paragraph is None:
        paragraph = document.add_paragraph("")

    try:
        if str_node == "\n" or str_node == "\xa0":
            return

        run = paragraph.add_run(str_node)

        if (parent.name == "strong") or ("bold" in style):
            run.font.bold = True
        if (parent.name == "em") or ("italic" in style):
            run.font.italic = True
        if "small" in style:
            run.font.size = Pt(8)
        if "underline" in style:
            run.font.underline = True
        if "left" in style:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        if "right" in style:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        if "center" in style:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if "justify" in style:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    except Exception:
        pass


def process_list(node, document):
    list_style = "List Bullet" if node.name == "ul" else "List Number"

    for child in node.children:
        if child.name == "li":
            paragraph = document.add_paragraph("", style=list_style)
            process_list_item(child, document, node, paragraph)


def process_list_item(node, document, parent, paragraph):
    try:
        remove_newlines(node)

        if node.contents:
            for child in node.children:
                if child.name == "math":
                    process_math(child, document, paragraph)
                    continue

                process_list_item(child, document, node, paragraph)

    except AttributeError:
        process_string(node, document, parent, paragraph)


def process_image(img_node, document, paragraph, styles=[]):
    if paragraph and paragraph.text == "":
        delete_paragraph(paragraph)

    url = img_node["src"]
    image = download_image(url)

    try:
        if should_use_width(img_node):
            size = convert_image_size(img_node["width"])
            document.add_picture(image, width=Mm(size))
        else:
            size = convert_image_size(img_node["height"])
            document.add_picture(image, height=Mm(size))

        if should_be_centered(img_node, styles):
            try:
                picture_paragraph = document.paragraphs[-1]
            except AttributeError:
                picture_paragraph = document._parent
            picture_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    except UnrecognizedImageError:
        document.add_paragraph(
                "### Could not include image from url: %s" % (url, )
        )


def process_math(node, document, paragraph=None):
    import re
    from django.conf import settings

    project_path = re.sub(r"/main$", "%s", settings.BASE_DIR)
    mml2omml_stylesheet_path = project_path % "/export/MML2OMML.XSL"
    element_tree = etree.fromstring(str(node))
    xslt = etree.parse(mml2omml_stylesheet_path)
    transform = etree.XSLT(xslt)
    new_dom = transform(element_tree)

    if paragraph is None:
        paragraph = document.add_paragraph()

    paragraph._element.append(new_dom.getroot())


def should_be_centered(img_node, styles=[]):
    centered = False
    try:
        style_str = img_node["style"]
        centered = ("margin-left: auto" in style_str) and (
            "margin-right: auto" in style_str
        )
    except:
        pass

    if "center" in styles:
        centered = True

    return centered


def convert_image_size(orginal_size):
    size_in_px = float(orginal_size)
    if size_in_px > MAX_PICTURE_SIZE:
        size_in_px = MAX_PICTURE_SIZE
    document_size = round(A4_PAGE_WIDHT - (DEFAULT_MARGINS * 2), 1)
    size_in_mn = round((size_in_px * document_size) / MAX_PICTURE_SIZE)
    return size_in_mn


def should_use_width(img_node):
    try:
        width = 0
        height = 0
        width = int(img_node["width"])
        height = int(img_node["height"])
    except:
        pass

    use_width = True
    if height > width:
        use_width = False

    return use_width


def process_table(document, table_node):
    all_rows = table_node.find_all("tr")
    rows_count = len(all_rows)
    cells_count = list(map(lambda x: len(x.find_all(["th", "td"])), all_rows))
    max_columns = max(cells_count)
    table = document.add_table(rows=rows_count, cols=max_columns)

    current_row = 0
    for html_row in all_rows:
        row_cells = table.rows[current_row].cells

        current_cell_index = 0
        for html_cell in html_row.find_all(["th", "td"]):
            table_cell = row_cells[current_cell_index]
            process_table_cell(document, html_cell, table_cell, [], True)
            if html_cell.get("colspan") is not None:
                index_to_merge = current_cell_index + int(html_cell["colspan"]) - 1
                cell_to_merge = row_cells[index_to_merge]
                table_cell.merge(cell_to_merge)
            current_cell_index = current_cell_index + 1
        current_row = current_row + 1


def process_table_cell(document, node, cell, styles=[], skip_newlines=False):
    # still not considering new paragraphs inside table_cell
    # https://github.com/python-openxml/python-docx/issues/216
    if node.get("style") is not None:
        styles = styles + process_style(node["style"])

    if not cell.paragraphs:
        paragraph = cell.add_paragraph()
    else:
        paragraph = cell.paragraphs[-1]

    for child in node.children:
        if isinstance(child, NavigableString):
            if skip_newlines and (not child.strip()):
                continue

            process_string(child, document, node, paragraph, styles)
        elif child.name == "math":
            process_math(child, document, paragraph)
        elif child.name == "img":
            paragraph = cell.add_paragraph()
            run = paragraph.add_run()
            process_image(child, run, None, styles)
            cell.add_paragraph()
        else:
            process_table_cell(document, child, cell, styles)


def process_style(style_str=""):
    import re

    default_size = 11
    styles = []

    if style_str is None:
        return styles

    sizeMatch = re.search("font-size: (\d+)pt", style_str)
    if sizeMatch:
        size = int(sizeMatch.group(1))
        if size < default_size:
            styles.append("small")

    weightMatch = re.search("font-weight: (700|800|900|bold(er)?)", style_str)
    if weightMatch:
        styles.append("bold")

    styleMatch = re.search("font-style: (oblique|italic)", style_str)
    if styleMatch:
        styles.append("italic")

    underMatch = re.search("text-decoration: underline", style_str)
    if underMatch:
        styles.append("underline")

    alignMatch = re.search("text-align: (left|right|center|justify)", style_str)
    if alignMatch:
        styles.append(alignMatch.group(1))

    return styles


def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None


def remove_newlines(node):
    all_elements = node.contents
    for element in all_elements:
        if isinstance(element, NavigableString):
            if element.strip() == "":
                element.extract()


def download_image(url):
    response = requests.get(url, stream=True)
    image = io.BytesIO(response.content)
    return image


def remove_table_border(table):
    from docx.oxml.shared import OxmlElement  # Necessary Import
    from docx.oxml.ns import qn

    # https://github.com/python-openxml/python-docx/issues/433
    tbl = table._tbl  # get xml element in table
    for cell in tbl.iter_tcs():
        tcPr = cell.tcPr  # get tcPr element, in which we can define style of borders
        tcBorders = OxmlElement("w:tcBorders")
        borders = []

        for element in ["w:top", "w:left", "w:bottom", "w:right"]:
            el = OxmlElement(element)
            el.set(qn("w:val"), "nil")
            borders.append(el)

        for b in borders:
            tcBorders.append(b)

        tcPr.append(tcBorders)
